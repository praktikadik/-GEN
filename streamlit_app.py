import streamlit as st
from docx import Document
import re
import io
from datetime import datetime

# --- ΡΥΘΜΙΣΕΙΣ ΣΕΛΙΔΑΣ ---
st.set_page_config(page_title="ΝΕΑ ΑΔΕΙΑ Online", page_icon="📝")

# --- ΣΥΣΤΗΜΑ password ---
def check_password():
    """Επιστρέφει True αν ο χρήστης έδωσε τον σωστό κωδικό."""
    if "password_correct" not in st.session_state:
        # Αρχικοποίηση
        st.session_state["password_correct"] = False

    if st.session_state["password_correct"]:
        return True

    # Εμφάνιση πεδίου κωδικού
    st.title("🔒 Πρόσβαση Περιορισμένη")
    password = st.text_input("Εισάγετε τον κωδικό πρόσβασης:", type="password")
    
    # Εδώ ορίζεις τον κωδικό σου (π.χ. 1234)
    if st.button("Είσοδος"):
        if password == st.secrets["MY_PASSWORD"]: # <--- ΑΛΛΑΞΕ ΤΟΝ ΚΩΔΙΚΟ ΕΔΩ
            st.session_state["password_correct"] = True
            st.rerun()
        else:
            st.error("❌ Λάθος κωδικός!")
    return False

# Αν ο κωδικός είναι σωστός, τρέξε την εφαρμογή
if check_password():
    st.title("📝 Νέα Άδεια")
    st.markdown("Συμπληρώστε τα πεδία παρακάτω για να δημιουργήσετε το έγγραφό σας.")

    # --- ΡΥΘΜΙΣΕΙΣ ΑΔΕΙΩΝ ---
    TYPES_OF_LEAVE = {
        "kanoniki": {
            "label": "Κανονική Άδεια",
            "template": "KANONIKI_ADEIA.docx",
            "output_prefix": "ΚΑΝΟΝΙΚΗ ΑΔΕΙΑ"
        },
        "mikri": {
            "label": "Άδεια Μικρής Διάρκειας",
            "template": "ADEIA_MIKRIS_DIARKEIAS.docx",
            "output_prefix": "ΑΔΕΙΑ ΜΙΚΡΗΣ ΔΙΑΡΚΕΙΑΣ"
        }
    }

    choice = st.radio("Επιλέξτε τύπο άδειας:", 
                      options=list(TYPES_OF_LEAVE.keys()), 
                      format_func=lambda x: TYPES_OF_LEAVE[x]["label"],
                      horizontal=True)

    config = TYPES_OF_LEAVE[choice]

    try:
        doc = Document(config["template"])
        
        text_blocks = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text_blocks.append(cell.text)
        
        full_text = "\n".join(text_blocks)
        placeholders = sorted(list(set(re.findall(r"\{\{(.*?)\}\}", full_text))))

        st.subheader(f"Στοιχεία για: {config['label']}")
        user_inputs = {}
        
        cols = st.columns(2)
        for i, placeholder in enumerate(placeholders):
            col = cols[i % 2]
            user_inputs[placeholder] = col.text_input(f"Εισάγετε {placeholder}:", key=placeholder)

        if st.button("ΠΡΟΕΤΟΙΜΑΣΙΑ ΕΓΓΡΑΦΟΥ", type="primary", use_container_width=True):
            replacements = {f"{{{{{k}}}}}": v for k, v in user_inputs.items()}
            
            def replace_in_text(text):
                for k, v in replacements.items():
                    if k in text: text = text.replace(k, v)
                return text

            for p in doc.paragraphs:
                for run in p.runs: run.text = replace_in_text(run.text)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs: run.text = replace_in_text(run.text)

            bio = io.BytesIO()
            doc.save(bio)
            
            today_str = datetime.now().strftime("%d-%m-%Y")
            file_name = f"{config['output_prefix']} {today_str}.docx"

            st.success("✅ Το έγγραφο είναι έτοιμο!")
            st.download_button(
                label="📥 ΛΗΨΗ ΑΡΧΕΙΟΥ (WORD)",
                data=bio.getvalue(),
                file_name=file_name,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

    except FileNotFoundError:
        st.error(f"⚠️ Το αρχείο {config['template']} δεν βρέθηκε.")
    except Exception as e:
        st.error(f"Σφάλμα: {e}")
